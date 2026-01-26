import numpy as np
import pandas as pd
import matplotlib.pyplot as plt

# Load both CSV files

def plot_overall_comparison(df1, df2, offset1, offset2):

    # Remove the rows where ratio is -1
    df1 = df1[df1['ratio'] != -1]
    df2 = df2[df2['ratio'] != -1]
    df1 = df1[(df1['offset']==offset1)]
    df2 = df2[(df2['offset']==offset2)]

    fingerprints = ['ECFP4', 'MACCS', 'RDKIT']

    means1 = []
    stds1 = []
    means2 = []
    stds2 = []

    for fp in fingerprints:
        # Dataset 1
        if fp == 'ECFP4':
            ratios1 = df1[df1['suffix'].str.endswith('ECFP')]['ratio']
        else:
            ratios1 = df1[df1['suffix'].str.endswith(fp)]['ratio']
        means1.append(ratios1.mean())
        stds1.append(ratios1.std())

        # Dataset 2
        if fp == 'ECFP4':
            ratios2 = df2[df2['suffix'].str.endswith('ECFP')]['ratio']
        else:
            ratios2 = df2[df2['suffix'].str.endswith(fp)]['ratio']
        means2.append(ratios2.mean())
        stds2.append(ratios2.std())

    plt.figure(figsize=(10, 6))

    x = np.arange(len(fingerprints))
    width = 0.35
    
    df1_recursive = df1['recursive'].unique()
    df2_recursive = df2['recursive'].unique()

    if len(df1_recursive) != 1 or len(df2_recursive) != 1:
        raise ValueError("DataFrames must have a single unique value for 'recursive' column.")
    df1_recursive = df1_recursive[0]
    df2_recursive = df2_recursive[0]
    if df1_recursive==False:
        label1='Non-Recursive'
    else:
        label1='Recursive'

    if df2_recursive==False:
        label2='Non-Recursive'
    else:
        label2='Recursive'

    bars1 = plt.bar(x - width/2, means1, width, 
                    color=['tab:blue', 'tab:orange', 'tab:green'], 
                alpha=0.8, label=f'{label1}, Offset={offset1}') # TODO: Remove hardcoding

    bars2 = plt.bar(x + width/2, means2, width, 
                    color=['lightblue', 'peachpuff', 'lightgreen'], 
                alpha=0.8, label=f'{label2}, Offset={offset2}') # TODO: Remove hardcoding

    plt.xlabel('Fingerprint', fontsize=12)
    plt.ylabel('Mean Ratio', fontsize=12)
    plt.title('Mean Ratio by Fingerprint Type', fontsize=14)
    plt.xticks(x, fingerprints)
    plt.legend()
    plt.tight_layout()
    plt.savefig(f'plots/overall_comparison_{label1}_{offset1}_{label2}_{offset2}.png')
    plt.show()
    plt.close()

def plot_by_threshold(df, offset, output_file):

    df = df[df['ratio'] != -1]
    df = df[(df['offset']==offset)]
    fingerprints = ['ECFP4', 'MACCS', 'RDKIT']
    colors = ['tab:blue', 'tab:orange', 'tab:green']

    plt.figure(figsize=(8,6))

    df = df.sort_values(by='threshold')

    for fp, color in zip(fingerprints, colors):

        means=[]
        stds=[]
        thresholds=[]

        for th in sorted(df['threshold'].unique()):
            df_th = df[df['threshold'] == th]
            ratios = df_th[df_th['suffix'].str.endswith(fp)]['ratio']
            if not ratios.empty:
                means.append(ratios.mean())
                stds.append(ratios.std())
                thresholds.append(th)

        plt.plot(thresholds, means, '-o', label=fp, color=color, linewidth=2, marker='s')
    
    plt.title(f'Ratio vs Threshold for each Fingerprint', fontsize=14)
    plt.xlabel('Threshold', fontsize=12)
    plt.ylabel('Ratio', fontsize=12)
    plt.legend(title='Fingerprint')
    plt.grid(False)
    plt.tight_layout()
    plt.savefig(output_file)
    plt.show()
    plt.close()


def plot_comparison_by_threshold(df1, df2, offset1, offset2, zoom_min=None, zoom_max=None, y_min=None, y_max=None):

    df1 = df1[df1['ratio'] != -1]
    df2 = df2[df2['ratio'] != -1]
    df1 = df1[(df1['offset']==offset1)]
    df2 = df2[(df2['offset']==offset2)]
    df1_recursive = df1['recursive'].unique()
    df2_recursive = df2['recursive'].unique()

    if len(df1_recursive) != 1 or len(df2_recursive) != 1:
        raise ValueError("DataFrames must have a single unique value for 'recursive' column.")
    df1_recursive = df1_recursive[0]
    df2_recursive = df2_recursive[0]
    if df1_recursive==False:
        label1='Non-Recursive'
    else:
        label1='Recursive'

    if df2_recursive==False:
        label2='Non-Recursive'
    else:
        label2='Recursive'

    fingerprints = ['ECFP4', 'MACCS', 'RDKIT']
    colors = ['tab:blue', 'tab:orange', 'tab:green']

    # Collect all thresholds to determine appropriate range
    all_thresholds = sorted(set(df1['Threshold'].unique()) | set(df2['Threshold'].unique()))
    
    if zoom_min is None or zoom_max is None:
        threshold_min, threshold_max = min(all_thresholds), max(all_thresholds)
        if zoom_min is None:
            # Check if 0.5 is in the range
            zoom_min = 0.5 if (0.5 >= threshold_min and 0.5 <= threshold_max) else threshold_min
        if zoom_max is None:
            # Check if 1.0 is in the range
            zoom_max = 1.0 if (1.0 >= threshold_min and 1.0 <= threshold_max) else threshold_max

    plt.figure(figsize=(10, 6))

    all_means_total = []

    # Keep all thresholds in the selected zoom range so x-axis shows true values
    thresholds_in_range = [t for t in all_thresholds if zoom_min <= t <= zoom_max]

    for fp, color in zip(fingerprints, colors):
        means1=[]
        means2=[]

        for th in thresholds_in_range:
            df1_th = df1[df1['Threshold'] == th]
            if fp == 'ECFP4':
                ratios1 = df1_th[df1_th['suffix'].str.endswith('ECFP')]['ratio']
            else:
                ratios1 = df1_th[df1_th['suffix'].str.endswith(fp)]['ratio']

            df2_th = df2[df2['Threshold'] == th]
            if fp == 'ECFP4':
                ratios2 = df2_th[df2_th['suffix'].str.endswith('ECFP')]['ratio']
            else:
                ratios2 = df2_th[df2_th['suffix'].str.endswith(fp)]['ratio']

            # Keep all thresholds in range, using NaN for missing data to maintain spacing
            if not ratios1.empty:
                means1.append(ratios1.mean())
            else:
                means1.append(np.nan)

            if not ratios2.empty:
                means2.append(ratios2.mean())
            else:
                means2.append(np.nan)

        plt.plot(thresholds_in_range, means1, '-o', label=f'{fp} {label1} Offset={offset1}', color=color, linewidth=2, marker='s', clip_on=False)
        plt.plot(thresholds_in_range, means2, '--o', label=f'{fp} {label2} Offset={offset2}', color=color, linewidth=2, marker='o', clip_on=False)
        
        # Collect all data for y-axis limit calculation
        all_means_total.extend([m for m in means1 if not np.isnan(m)])
        all_means_total.extend([m for m in means2 if not np.isnan(m)])

    plt.title('Ratio vs Threshold Comparison', fontsize=14)
    plt.xlabel('Threshold', fontsize=12)
    plt.ylabel('Ratio', fontsize=12)
    
    # Set x-tick labels to show true threshold values with full precision
    if thresholds_in_range:
        plt.xticks(thresholds_in_range, [f'{t:.2f}' for t in thresholds_in_range], rotation=0)
        x_margin = (max(thresholds_in_range) - min(thresholds_in_range)) * 0.05
        plt.xlim(min(thresholds_in_range) - x_margin, max(thresholds_in_range) + x_margin)
    
    # Set y-axis limits naturally (matplotlib will auto-scale)
    if y_min is None:
        y_min = min(all_means_total) - 0.05 if all_means_total else 0.6
    plt.ylim(ymin=y_min)
    
    plt.legend(title='Fingerprint and Offset')
    plt.grid(False)
    plt.tight_layout()
    plt.savefig('plots/comparison_by_threshold_'+label1+'_'+str(offset1)+'_'+label2+'_'+str(offset2)+'.png')
    plt.show()
    plt.close()
    
    # Generate Word document with statistics (using the selected threshold range)
    from docx import Document
    import os
    
    doc = Document()
    doc.add_heading('Comparison by Threshold - Statistical Summary', 0)
    doc.add_paragraph(f'Threshold Range: {zoom_min:.2f} to {zoom_max:.2f}')
    doc.add_paragraph()
    
    # Recreate the data for the Word document (with threshold range 0.5 to 1.0)
    doc.add_heading(f'{label1} (Offset={offset1})', level=1)
    

    for fp in fingerprints:
        doc.add_heading(f'{fp}', level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Threshold'
        hdr_cells[1].text = 'Mean'
        hdr_cells[2].text = 'Std Dev'
        means1=[]
        stds1=[]
        thresholds=[]
        for th in sorted(df1['Threshold'].unique()):
            df1_th = df1[df1['Threshold'] == th]
            if fp == 'ECFP4':
                ratios1 = df1_th[df1_th['suffix'].str.endswith('ECFP')]['ratio']
            else:
                ratios1 = df1_th[df1_th['suffix'].str.endswith(fp)]['ratio']
            if not ratios1.empty:
                means1.append(ratios1.mean())
                stds1.append(ratios1.std())
                thresholds.append(th)
        # Filter to the selected threshold range
        filtered_data = [(th, m, s) for th, m, s in zip(thresholds, means1, stds1) if zoom_min <= th <= zoom_max]
        for th, mean, std in filtered_data:
            row_cells = table.add_row().cells
            row_cells[0].text = f"{th:.2f}"
            row_cells[1].text = f"{mean:.6f}"
            row_cells[2].text = f"{std:.6f}"
        doc.add_paragraph()
    
    doc.add_heading(f'{label2} (Offset={offset2})', level=1)
    

    for fp in fingerprints:
        doc.add_heading(f'{fp}', level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Threshold'
        hdr_cells[1].text = 'Mean'
        hdr_cells[2].text = 'Std Dev'
        means2=[]
        stds2=[]
        thresholds=[]
        for th in sorted(df2['Threshold'].unique()):
            df2_th = df2[df2['Threshold'] == th]
            if fp == 'ECFP4':
                ratios2 = df2_th[df2_th['suffix'].str.endswith('ECFP')]['ratio']
            else:
                ratios2 = df2_th[df2_th['suffix'].str.endswith(fp)]['ratio']
            if not ratios2.empty:
                means2.append(ratios2.mean())
                stds2.append(ratios2.std())
                thresholds.append(th)
        # Filter to the selected threshold range
        filtered_data = [(th, m, s) for th, m, s in zip(thresholds, means2, stds2) if zoom_min <= th <= zoom_max]
        for th, mean, std in filtered_data:
            row_cells = table.add_row().cells
            row_cells[0].text = f"{th:.2f}"
            row_cells[1].text = f"{mean:.6f}"
            row_cells[2].text = f"{std:.6f}"
        doc.add_paragraph()
    
    # Save the document
    os.makedirs('results', exist_ok=True)
    doc.save(f'results/comparison_by_threshold_{label1}_{offset1}_{label2}_{offset2}.docx')
    print(f"Word document generated: results/comparison_by_threshold_{label1}_{offset1}_{label2}_{offset2}.docx")

def plot_num_molecule(df, fps=['ECFP4', 'MACCS', 'RDKIT']): #Offset is not used now
    

    colors=plt.cm.viridis

    for fp in fps:
        plt.figure(figsize=(10,6))
        # Map ECFP4 to ECFP for filtering
        fp_filter = 'ECFP' if fp == 'ECFP4' else fp
        df_fp = df[df['fingerprint_type']==fp_filter]
        thresholds = sorted(df_fp['threshold'].unique())
        color_map = {th: colors(i / (len(thresholds)-1)) for i, th in enumerate(thresholds)}
        for th in thresholds:
            df_th = df_fp[df_fp['threshold'] == th]
            x = df_th['cluster_index']
            plt.plot(x, df_th['cliff_nmols'], '-o', color=color_map[th], label=f'Cliff th={th:.2f}' , alpha=0.8)
            plt.plot(x, df_th['smooth_nmols'], '--o', color=color_map[th], label=f'Smooth th={th:.2f}' , alpha=0.8)
        plt.title(f'{fp}: Cluster Size vs Cluster Number')
        plt.xlabel('Cluster Number')
        plt.ylabel('Number of Molecules')
        # Custom legend for only one threshold (to avoid duplicate labels)
        handles, labels = plt.gca().get_legend_handles_labels()
        by_label = dict(zip(labels, handles))
        plt.legend(by_label.values(), by_label.keys(), fontsize=10)
        plt.grid(False)
        plt.tight_layout()
        plt.savefig(f'Plots/Smooth/cluster_size_vs_number_{fp}.png', dpi=300)
        plt.show()

def plot_prop_std(df, fps=['ECFP4', 'MACCS', 'RDKIT']): #Offset is not used now
        colors=plt.cm.viridis

        for fp in fps:
          plt.figure(figsize=(10, 6))
          fp_filter = 'ECFP' if fp == 'ECFP4' else fp
          df_fp = df[df['fingerprint_type']==fp_filter]
          thresholds = sorted(df_fp['threshold'].unique())
          color_map = {th: colors(i / (len(thresholds)-1)) for i, th in enumerate(thresholds)}
          for th in thresholds:
              df_th = df_fp[df_fp['threshold'] == th]
              x = df_th['cluster_index']
              # Cliff std (solid)
              plt.plot(
                  x, df_th['cliff_p_std'],
                  '-o', color=color_map[th], label=f'Cliff th={th:.2f}' , alpha=0.8
              )
              # Smooth std (dashed)
              plt.plot(
                  x, df_th['smooth_p_std'],
                  '--o', color=color_map[th], label=f'Smooth th={th:.2f}' , alpha=0.8
              )
          plt.title(f'{fp}: Property Std vs Cluster Number')
          plt.xlabel('Cluster Number')
          plt.ylabel('Property Std')
          plt.ylim(0, 2.5)
          # Custom legend for only one threshold (to avoid duplicate labels)
          handles, labels = plt.gca().get_legend_handles_labels()
          by_label = dict(zip(labels, handles))
          plt.legend(by_label.values(), by_label.keys(), fontsize=10)
          plt.grid(False)
          plt.tight_layout()
          plt.savefig(f'Plots/Smooth/cluster_property_std_vs_number_{fp}.png', dpi=300)
          plt.show()


def all_comparisons(df_no_recur, df_recur):
    """
    Plot all comparisons in a single plot:
    Order: Non-recursive Offset=0, Recursive Offset=0, Non-recursive Offset=0.3, Recursive Offset=0.3
    Also generates a Word document with mean and standard deviation statistics.
    """
    from docx import Document
    
    df_no_recur = df_no_recur[df_no_recur['ratio'] != -1]
    df_recur = df_recur[df_recur['ratio'] != -1]

    # Get data for all 4 combinations in order
    df_list = [
        df_no_recur[df_no_recur['offset'] == 0.0],    # Non-recursive, Offset=0
        df_recur[df_recur['offset'] == 0.0],           # Recursive, Offset=0
        df_no_recur[df_no_recur['offset'] == 0.3],    # Non-recursive, Offset=0.3
        df_recur[df_recur['offset'] == 0.3]            # Recursive, Offset=0.3
    ]
    
    labels_list = [
        'Non-Recursive, Offset=0',
        'Recursive, Offset=0',
        'Non-Recursive, Offset=0.3',
        'Recursive, Offset=0.3'
    ]
    
    fingerprints = ['ECFP4', 'MACCS', 'RDKIT']
    
    # Collect means and stds for all combinations
    all_means = [[], [], [], []]
    all_stds = [[], [], [], []]
    
    for idx, df in enumerate(df_list):
        for fp in fingerprints:
            if fp == 'ECFP4':
                ratios = df[df['suffix'].str.endswith('ECFP')]['ratio']
            else:
                ratios = df[df['suffix'].str.endswith(fp)]['ratio']
            all_means[idx].append(ratios.mean())
            all_stds[idx].append(ratios.std())
    
    # Create plots
    plt.figure(figsize=(15, 9))
    
    x = np.arange(len(fingerprints))
    width = 0.2
    
    # Define color schemes for each fingerprint type
    # Lighter shades for offset=0, darker shades for offset=0.3
    fp_colors = {
        'ECFP4': ['#aec7e8', '#7ba3d1', '#1f77b4', '#0d47a1'],    # Light to dark blues
        'MACCS': ['#ffbb78', '#ff9d3d', '#ff7f0e', '#cc5200'],   # Light to dark oranges
        'RDKIT': ['#98df8a', '#5bb56c', '#2ca02c', '#1a6b1f']    # Light to dark greens
    }
    
    # Create bars for each combination
    for idx in range(4):
        offset_pos = x + (idx - 1.5) * width
        colors_for_bars = [fp_colors[fp][idx] for fp in fingerprints]
        plt.bar(offset_pos, all_means[idx], width,
                color=colors_for_bars, alpha=0.9, label=labels_list[idx], clip_on=False)
    
    plt.xlabel('Fingerprint', fontsize=12)
    plt.ylabel('Mean Ratio', fontsize=12)
    
    plt.title('Mean Ratio by Fingerprint Type and Method', fontsize=14)
    plt.xticks(x, fingerprints)
    plt.legend()
    plt.tight_layout()
    plt.savefig('plots/all_comparisons_ordered.png')
    plt.show()
    plt.close()
    
    # Generate Word document with statistics
    doc = Document()
    doc.add_heading('All Comparisons - Statistical Summary', 0)

    # Create a table for each combination
    for idx, (label, means, stds) in enumerate(zip(labels_list, all_means, all_stds)):
        doc.add_heading(label, level=1)
        
        # Create table with headers
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Fingerprint'
        hdr_cells[1].text = 'Mean'
        hdr_cells[2].text = 'Std Dev'
        
        # Data rows
        for fp, mean, std in zip(fingerprints, means, stds):
            row_cells = table.add_row().cells
            row_cells[0].text = fp
            row_cells[1].text = f"{mean:.6f}"
            row_cells[2].text = f"{std:.6f}"
        
        doc.add_paragraph()  # Add spacing between tables

    # Save the document
    doc.save('results/all_comparisons_statistics.docx')
    print("Word document generated: results/all_comparisons_statistics.docx")


